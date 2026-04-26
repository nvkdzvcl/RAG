"""API tests for document upload and processing status routes."""

from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace

import pytest
from fastapi.testclient import TestClient

from app.core.config import get_settings
from app.main import create_app


@pytest.fixture
def isolated_client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> tuple[TestClient, Path]:
    data_dir = tmp_path / "data"
    corpus_dir = tmp_path / "corpus"
    corpus_dir.mkdir(parents=True, exist_ok=True)
    (corpus_dir / "seed.txt").write_text(
        "This seeded corpus text should be replaced by uploaded indexes when available.",
        encoding="utf-8",
    )

    monkeypatch.setenv("DATA_DIR", str(data_dir))
    monkeypatch.setenv("INDEX_DIR", str(data_dir / "indexes"))
    monkeypatch.setenv("CORPUS_DIR", str(corpus_dir))
    get_settings.cache_clear()

    app = create_app()
    client = TestClient(app)
    try:
        yield client, data_dir
    finally:
        client.close()
        get_settings.cache_clear()


@pytest.fixture
def isolated_client_ocr_enabled(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> tuple[TestClient, Path]:
    data_dir = tmp_path / "data"
    corpus_dir = tmp_path / "corpus"
    corpus_dir.mkdir(parents=True, exist_ok=True)
    (corpus_dir / "seed.txt").write_text(
        "Seeded corpus for OCR-enabled upload fixture.",
        encoding="utf-8",
    )

    class FakePage:
        images = []

        @staticmethod
        def extract_text() -> str:
            return ""

        @staticmethod
        def extract_tables() -> list[list[list[str]]]:
            return []

    class FakePDF:
        def __init__(self) -> None:
            self.pages = [FakePage()]

        def __enter__(self) -> "FakePDF":
            return self

        def __exit__(self, exc_type, exc, tb) -> None:
            _ = exc_type
            _ = exc
            _ = tb

    monkeypatch.setenv("DATA_DIR", str(data_dir))
    monkeypatch.setenv("INDEX_DIR", str(data_dir / "indexes"))
    monkeypatch.setenv("CORPUS_DIR", str(corpus_dir))
    monkeypatch.setenv("OCR_ENABLED", "true")
    monkeypatch.setenv("OCR_LANGUAGE", "vie+eng")
    monkeypatch.setenv("OCR_MIN_TEXT_CHARS", "100")
    monkeypatch.setenv("OCR_RENDER_DPI", "216")
    monkeypatch.setenv("OCR_CONFIDENCE_THRESHOLD", "40")
    monkeypatch.setenv("TESSERACT_CMD", "")
    monkeypatch.setattr(
        "app.ingestion.parsers.pdf_parser.pdfplumber",
        SimpleNamespace(open=lambda _: FakePDF()),
    )
    monkeypatch.setattr("app.ingestion.parsers.pdf_parser.is_tesseract_available", lambda: True)
    monkeypatch.setattr(
        "app.ingestion.parsers.pdf_parser.ocr_pdf_page_with_pymupdf",
        lambda *args, **kwargs: "Nội dung OCR tiếng Việt: tokendebugocr-77",
    )
    get_settings.cache_clear()

    app = create_app()
    client = TestClient(app)
    try:
        yield client, data_dir
    finally:
        client.close()
        get_settings.cache_clear()


def test_upload_endpoint_returns_document_payload(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, data_dir = isolated_client

    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("uploaded.txt", b"Upload flow should build hybrid indexes from this chunk.", "text/plain")},
    )

    assert response.status_code == 201
    body = response.json()

    assert body["document_id"]
    assert body["id"] == body["document_id"]
    assert body["filename"] == "uploaded.txt"
    assert body["original_filename"] == "uploaded.txt"
    assert body["file_type"] == "txt"
    assert body["status"] == "ready"
    assert body["chunk_count"] is not None
    assert body["chunk_count"] > 0
    assert body["created_at"]
    assert body["uploaded_at"]

    raw_files = sorted((data_dir / "raw").glob("*"))
    assert raw_files
    assert any(path.name.endswith("_uploaded.txt") for path in raw_files)


def test_upload_response_reports_ocr_debug_metadata(
    isolated_client_ocr_enabled: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client_ocr_enabled

    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("testocr.pdf", b"%PDF-1.4 fake scan", "application/pdf")},
    )

    assert response.status_code == 201
    body = response.json()
    assert body["status"] == "ready"
    assert "total_blocks" in body
    assert "text_blocks" in body
    assert "table_blocks" in body
    assert "image_blocks" in body
    assert "ocr_blocks" in body
    assert "total_chunks" in body
    assert "ocr_chunks" in body
    assert body["total_blocks"] >= 1
    assert body["text_blocks"] >= 1
    assert body["ocr_blocks"] >= 1
    assert body["total_chunks"] >= 1
    assert body["ocr_chunks"] >= 1


def test_ocr_chunk_is_indexed_retrievable_and_marked_in_sources(
    isolated_client_ocr_enabled: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client_ocr_enabled

    upload_response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("testocr.pdf", b"%PDF-1.4 fake scan", "application/pdf")},
    )
    assert upload_response.status_code == 201

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "tokendebugocr-77 là gì?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_response.status_code == 200
    body = query_response.json()
    assert body["status"] != "insufficient_evidence"
    assert body["trace"][0]["index_source"] == "uploaded"
    assert any(citation.get("block_type") == "ocr_text" for citation in body["citations"])
    context_steps = [step for step in body["trace"] if step.get("step") == "context_select"]
    assert context_steps
    selected_docs = context_steps[0].get("docs", [])
    assert any(doc.get("block_type") == "ocr_text" for doc in selected_docs)


def test_upload_endpoint_accepts_docx(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    from docx import Document

    buffer = BytesIO()
    doc = Document()
    doc.add_heading("Báo cáo", level=1)
    doc.add_paragraph("Tài liệu DOCX có bảng và tiếng Việt.")
    doc.save(buffer)
    buffer.seek(0)

    response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "report.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 201
    body = response.json()
    assert body["filename"] == "report.docx"
    assert body["status"] == "ready"


def test_document_status_endpoint_returns_saved_state(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "guide.md",
                b"# Guide\n\nRuntime indexing should move this document to ready state.",
                "text/markdown",
            )
        },
    )
    assert upload_response.status_code == 201
    document_id = upload_response.json()["document_id"]

    response = client.get(f"/api/v1/documents/{document_id}/status")
    assert response.status_code == 200

    body = response.json()
    assert body["document_id"] == document_id
    assert body["filename"] == "guide.md"
    assert body["status"] == "ready"
    assert body["chunk_count"] is not None

    listed = client.get("/api/v1/documents")
    assert listed.status_code == 200
    listed_body = listed.json()
    assert "documents" in listed_body
    assert any(item["document_id"] == document_id for item in listed_body["documents"])


def test_query_uses_uploaded_document_chunks_when_available(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "knowledge.txt",
                b"Blue-ocean sentinel token: nimbus-42. This phrase is unique to uploaded docs.",
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 201
    uploaded_document_id = upload_response.json()["document_id"]

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "What does nimbus-42 refer to?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_response.status_code == 200

    body = query_response.json()
    citations = body["citations"]
    assert citations
    assert any(citation["doc_id"] == uploaded_document_id for citation in citations)
    assert body["trace"][0]["index_source"] == "uploaded"


def test_query_trace_and_citations_include_uploaded_file_metadata(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "meta-source.txt",
                b"Unique source metadata token: metadata-source-551.",
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 201
    uploaded_document_id = upload_response.json()["document_id"]

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "metadata-source-551 la gi?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_response.status_code == 200
    body = query_response.json()

    context_steps = [step for step in body["trace"] if step.get("step") == "context_select"]
    assert context_steps
    docs = context_steps[0]["docs"]
    assert docs
    first_doc = docs[0]
    assert first_doc["doc_id"] == uploaded_document_id
    assert first_doc.get("file_name") == "meta-source.txt"
    assert first_doc.get("file_type") == "txt"
    assert "uploaded_at" in first_doc or "created_at" in first_doc
    assert "page" in first_doc
    assert "block_type" in first_doc
    assert "ocr" in first_doc

    citations = body["citations"]
    assert citations
    assert any(citation.get("file_name") == "meta-source.txt" for citation in citations)
    assert any(citation.get("file_type") == "txt" for citation in citations)
    assert any(citation.get("doc_id") == uploaded_document_id for citation in citations)
    first_citation = citations[0]
    assert first_citation.get("source_id") == first_citation.get("chunk_id")
    assert first_citation.get("filename") == "meta-source.txt"
    assert "metadata-source-551" in first_citation.get("content", "")
    assert "metadata-source-551" in first_citation.get("text", "")
    assert "metadata-source-551" in first_citation.get("snippet", "")
    assert isinstance(first_citation.get("score"), int | float)


def test_query_filter_doc_id_a_returns_only_document_a(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_a = client.post(
        "/api/v1/documents/upload",
        files={"file": ("filter-a.txt", b"shared-filter-token-991 alpha-doc-only-a", "text/plain")},
    )
    upload_b = client.post(
        "/api/v1/documents/upload",
        files={"file": ("filter-b.txt", b"shared-filter-token-991 beta-doc-only-b", "text/plain")},
    )
    assert upload_a.status_code == 201
    assert upload_b.status_code == 201
    doc_a = upload_a.json()["document_id"]
    doc_b = upload_b.json()["document_id"]

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "shared-filter-token-991 nằm ở đâu?",
            "mode": "standard",
            "chat_history": [],
            "doc_ids": [doc_a],
        },
    )
    assert query_response.status_code == 200
    body = query_response.json()

    assert body["trace"][0]["index_source"] == "uploaded"
    assert body["trace"][0]["applied_filters"]["doc_ids"] == [doc_a]
    assert body["trace"][0]["candidate_count_before_filter"] >= body["trace"][0]["candidate_count_after_filter"]
    citations = body["citations"]
    assert citations
    assert all(citation["doc_id"] == doc_a for citation in citations)
    assert all(citation["doc_id"] != doc_b for citation in citations)


def test_query_filter_doc_id_b_returns_only_document_b(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_a = client.post(
        "/api/v1/documents/upload",
        files={"file": ("filter-a.txt", b"shared-filter-token-776 alpha-doc-only-a", "text/plain")},
    )
    upload_b = client.post(
        "/api/v1/documents/upload",
        files={"file": ("filter-b.txt", b"shared-filter-token-776 beta-doc-only-b", "text/plain")},
    )
    assert upload_a.status_code == 201
    assert upload_b.status_code == 201
    doc_a = upload_a.json()["document_id"]
    doc_b = upload_b.json()["document_id"]

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "shared-filter-token-776 nằm ở đâu?",
            "mode": "standard",
            "chat_history": [],
            "doc_ids": [doc_b],
        },
    )
    assert query_response.status_code == 200
    body = query_response.json()

    assert body["trace"][0]["index_source"] == "uploaded"
    assert body["trace"][0]["applied_filters"]["doc_ids"] == [doc_b]
    citations = body["citations"]
    assert citations
    assert all(citation["doc_id"] == doc_b for citation in citations)
    assert all(citation["doc_id"] != doc_a for citation in citations)


def test_query_filter_by_filename_returns_expected_document(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_a = client.post(
        "/api/v1/documents/upload",
        files={"file": ("finance-a.txt", b"shared-filename-filter-311 in file a", "text/plain")},
    )
    upload_b = client.post(
        "/api/v1/documents/upload",
        files={"file": ("finance-b.txt", b"shared-filename-filter-311 in file b", "text/plain")},
    )
    assert upload_a.status_code == 201
    assert upload_b.status_code == 201
    doc_a = upload_a.json()["document_id"]
    doc_b = upload_b.json()["document_id"]

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "shared-filename-filter-311 la gi?",
            "mode": "standard",
            "chat_history": [],
            "filenames": ["finance-b.txt"],
        },
    )
    assert query_response.status_code == 200
    body = query_response.json()

    assert body["trace"][0]["applied_filters"]["filenames"] == ["finance-b.txt"]
    citations = body["citations"]
    assert citations
    assert all(citation["doc_id"] == doc_b for citation in citations)
    assert all(citation["doc_id"] != doc_a for citation in citations)


def test_deleted_document_cannot_be_retrieved_even_with_doc_id_filter(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_a = client.post(
        "/api/v1/documents/upload",
        files={"file": ("stale-a.txt", b"stale-doc-token-551 only in deleted doc", "text/plain")},
    )
    upload_b = client.post(
        "/api/v1/documents/upload",
        files={"file": ("stale-b.txt", b"live-doc-token-662", "text/plain")},
    )
    assert upload_a.status_code == 201
    assert upload_b.status_code == 201
    doc_a = upload_a.json()["document_id"]

    deleted = client.delete(f"/api/v1/documents/{doc_a}")
    assert deleted.status_code == 200

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "stale-doc-token-551 là gì?",
            "mode": "standard",
            "chat_history": [],
            "doc_ids": [doc_a],
        },
    )
    assert query_response.status_code == 200
    body = query_response.json()

    assert body["trace"][0]["index_source"] == "uploaded"
    assert body["trace"][0]["candidate_count_after_filter"] == 0
    assert body["citations"] == []
    assert body["status"] == "insufficient_evidence"


def test_unmatched_query_filter_does_not_fallback_to_unrelated_seeded_corpus(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "What is in the seeded corpus?",
            "mode": "standard",
            "chat_history": [],
            "doc_ids": ["missing-uploaded-doc-id"],
        },
    )
    assert query_response.status_code == 200
    body = query_response.json()

    retrieve_step = body["trace"][0]
    assert retrieve_step["index_source"] == "seeded"
    assert retrieve_step["applied_filters"]["doc_ids"] == ["missing-uploaded-doc-id"]
    assert retrieve_step["candidate_count_before_filter"] == 0
    assert retrieve_step["candidate_count_after_filter"] == 0
    assert retrieve_step["count"] == 0
    assert body["citations"] == []
    assert body["status"] == "insufficient_evidence"


def test_query_without_filters_remains_backward_compatible(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "What is in the seeded corpus?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_response.status_code == 200
    body = query_response.json()

    retrieve_step = body["trace"][0]
    assert retrieve_step["index_source"] == "seeded"
    assert retrieve_step["applied_filters"] == {}
    assert retrieve_step["count"] >= 1
    assert retrieve_step["candidate_count_before_filter"] == retrieve_step["count"]
    assert retrieve_step["candidate_count_after_filter"] == retrieve_step["count"]


def test_advanced_query_respects_doc_id_filter(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_a = client.post(
        "/api/v1/documents/upload",
        files={"file": ("adv-filter-a.txt", b"shared-adv-filter-121 alpha-doc-a", "text/plain")},
    )
    upload_b = client.post(
        "/api/v1/documents/upload",
        files={"file": ("adv-filter-b.txt", b"shared-adv-filter-121 beta-doc-b", "text/plain")},
    )
    assert upload_a.status_code == 201
    assert upload_b.status_code == 201
    doc_a = upload_a.json()["document_id"]
    doc_b = upload_b.json()["document_id"]

    response = client.post(
        "/api/v1/query",
        json={
            "query": "force retrieval shared-adv-filter-121 là gì?",
            "mode": "advanced",
            "chat_history": [],
            "doc_ids": [doc_b],
        },
    )
    assert response.status_code == 200
    body = response.json()

    loop_steps = [step for step in body["trace"] if step.get("step") == "loop"]
    assert loop_steps
    loop = loop_steps[0]
    assert loop["applied_filters"]["doc_ids"] == [doc_b]
    assert loop["candidate_count_before_filter"] >= loop["candidate_count_after_filter"]
    assert loop["retrieved_count"] >= 1

    selected_docs = loop.get("selected_context_docs", [])
    assert selected_docs
    assert all(doc["doc_id"] == doc_b for doc in selected_docs)
    assert all(doc["doc_id"] != doc_a for doc in selected_docs)

    citations = body.get("citations", [])
    if citations:
        assert all(citation["doc_id"] == doc_b for citation in citations)
        assert all(citation["doc_id"] != doc_a for citation in citations)


def test_compare_query_passes_filters_to_both_standard_and_advanced_branches(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    filtered_doc_id = "missing-uploaded-doc-id"
    response = client.post(
        "/api/v1/query",
        json={
            "query": "force retrieval compare-filter-check",
            "mode": "compare",
            "chat_history": [],
            "doc_ids": [filtered_doc_id],
        },
    )
    assert response.status_code == 200
    body = response.json()

    standard_retrieve = body["standard"]["trace"][0]
    assert standard_retrieve["step"] == "retrieve"
    assert standard_retrieve["applied_filters"]["doc_ids"] == [filtered_doc_id]
    assert standard_retrieve["candidate_count_before_filter"] == 0
    assert standard_retrieve["candidate_count_after_filter"] == 0
    assert standard_retrieve["count"] == 0

    advanced_loop_steps = [step for step in body["advanced"]["trace"] if step.get("step") == "loop"]
    assert advanced_loop_steps
    advanced_loop = advanced_loop_steps[0]
    assert advanced_loop["applied_filters"]["doc_ids"] == [filtered_doc_id]
    assert advanced_loop["candidate_count_before_filter"] == 0
    assert advanced_loop["candidate_count_after_filter"] == 0
    assert advanced_loop["retrieved_count"] == 0


def test_query_uses_seeded_corpus_when_no_uploaded_documents_ready(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "What is in the seeded corpus?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_response.status_code == 200

    body = query_response.json()
    assert body["trace"]
    assert body["trace"][0]["step"] == "retrieve"
    assert body["trace"][0]["index_source"] == "seeded"


def test_standard_mode_retrieves_vietnamese_uploaded_chunks(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "vi-knowledge.txt",
                "Hệ thống truy hồi thông tin dùng cụm từ khóa độc nhất: baobietviet42.".encode("utf-8"),
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 201
    uploaded_document_id = upload_response.json()["document_id"]

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "baobietviet42 có ý nghĩa gì trong tài liệu?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_response.status_code == 200

    body = query_response.json()
    citations = body["citations"]
    assert citations
    assert any(citation["doc_id"] == uploaded_document_id for citation in citations)


def test_delete_all_documents_clears_registry_raw_files_and_runtime_indexes(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, data_dir = isolated_client

    first_upload = client.post(
        "/api/v1/documents/upload",
        files={"file": ("doc-a.txt", b"alpha token xoa-all-001", "text/plain")},
    )
    second_upload = client.post(
        "/api/v1/documents/upload",
        files={"file": ("doc-b.txt", b"beta token xoa-all-002", "text/plain")},
    )
    assert first_upload.status_code == 201
    assert second_upload.status_code == 201
    deleted_ids = {
        first_upload.json()["document_id"],
        second_upload.json()["document_id"],
    }

    index_dir = data_dir / "indexes"
    uploaded_artifacts = [
        index_dir / "uploaded_vector_index.json",
        index_dir / "uploaded_bm25_index.json",
        index_dir / "uploaded_index_manifest.json",
    ]
    assert all(path.exists() for path in uploaded_artifacts)

    listed_before = client.get("/api/v1/documents")
    assert listed_before.status_code == 200
    assert len(listed_before.json()["documents"]) == 2

    delete_response = client.delete("/api/v1/documents")
    assert delete_response.status_code == 200
    payload = delete_response.json()
    assert payload["status"] == "deleted"
    assert payload["deleted_documents"] == 2
    assert payload["deleted_files"] >= 2

    listed_after = client.get("/api/v1/documents")
    assert listed_after.status_code == 200
    assert listed_after.json()["documents"] == []

    registry_payload = json.loads((data_dir / "document_registry.json").read_text(encoding="utf-8"))
    assert registry_payload["documents"] == []

    raw_files = [path for path in (data_dir / "raw").rglob("*") if path.is_file()]
    assert raw_files == []
    assert all(not path.exists() for path in uploaded_artifacts)

    query_after_delete = client.post(
        "/api/v1/query",
        json={
            "query": "xoa-all-001 là gì?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_after_delete.status_code == 200
    query_body = query_after_delete.json()
    assert query_body["trace"][0]["index_source"] == "seeded"
    assert all(citation["doc_id"] not in deleted_ids for citation in query_body["citations"])


def test_delete_all_documents_when_empty_is_idempotent(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    response = client.delete("/api/v1/documents")
    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "deleted"
    assert payload["deleted_documents"] == 0
    assert payload["deleted_files"] == 0


def test_delete_single_document_rebuilds_runtime_from_remaining_uploaded_docs(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, data_dir = isolated_client

    first_upload = client.post(
        "/api/v1/documents/upload",
        files={"file": ("first.txt", b"first doc token remove-this-doc", "text/plain")},
    )
    second_upload = client.post(
        "/api/v1/documents/upload",
        files={"file": ("second.txt", b"second doc token keep-this-doc-999", "text/plain")},
    )
    assert first_upload.status_code == 201
    assert second_upload.status_code == 201
    first_id = first_upload.json()["document_id"]
    second_id = second_upload.json()["document_id"]

    delete_response = client.delete(f"/api/v1/documents/{first_id}")
    assert delete_response.status_code == 200
    payload = delete_response.json()
    assert payload["status"] == "deleted"
    assert payload["document_id"] == first_id
    assert payload["remaining_documents"] == 1
    assert payload["deleted_files"] >= 1

    listed = client.get("/api/v1/documents")
    assert listed.status_code == 200
    listed_ids = [item["document_id"] for item in listed.json()["documents"]]
    assert first_id not in listed_ids
    assert second_id in listed_ids

    raw_files = [path.name for path in (data_dir / "raw").rglob("*") if path.is_file()]
    assert len(raw_files) == 1

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "keep-this-doc-999 có trong tài liệu nào?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_response.status_code == 200
    body = query_response.json()
    assert body["trace"][0]["index_source"] == "uploaded"
    assert any(citation["doc_id"] == second_id for citation in body["citations"])
    assert all(citation["doc_id"] != first_id for citation in body["citations"])


def test_upload_then_delete_same_document_never_returns_deleted_doc(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload = client.post(
        "/api/v1/documents/upload",
        files={"file": ("single-delete.txt", b"single-delete-token-4477", "text/plain")},
    )
    assert upload.status_code == 201
    doc_id = upload.json()["document_id"]

    before_delete = client.post(
        "/api/v1/query",
        json={
            "query": "single-delete-token-4477 la gi?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert before_delete.status_code == 200
    assert any(citation["doc_id"] == doc_id for citation in before_delete.json()["citations"])

    deleted = client.delete(f"/api/v1/documents/{doc_id}")
    assert deleted.status_code == 200

    deleted_status = client.get(f"/api/v1/documents/{doc_id}/status")
    assert deleted_status.status_code == 404

    listed = client.get("/api/v1/documents")
    assert listed.status_code == 200
    assert all(item["document_id"] != doc_id for item in listed.json()["documents"])

    after_delete = client.post(
        "/api/v1/query",
        json={
            "query": "single-delete-token-4477 la gi?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert after_delete.status_code == 200
    body = after_delete.json()
    assert body["trace"][0]["index_source"] == "seeded"
    assert all(citation["doc_id"] != doc_id for citation in body["citations"])


def test_delete_missing_document_returns_404(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    response = client.delete("/api/v1/documents/missing-document-id")
    assert response.status_code == 404
    assert "Document not found" in response.json()["detail"]


def test_reindex_endpoint_updates_chunk_strategy_and_keeps_service_healthy(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "chunking.txt",
                ("Chunk strategy token. " * 120).encode("utf-8"),
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 201

    reindex = client.post(
        "/api/v1/documents/reindex",
        json={"chunk_size": 500, "chunk_overlap": 50},
    )
    assert reindex.status_code == 200
    body = reindex.json()
    assert body["status"] == "reindexed"
    assert body["chunk_size"] == 500
    assert body["chunk_overlap"] == 50
    assert body["reindexed_documents"] >= 1

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "Chunk strategy token là gì?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_response.status_code == 200
    query_body = query_response.json()
    retrieve_step = query_body["trace"][0]
    assert retrieve_step["step"] == "retrieve"
    assert retrieve_step["chunk_size"] == 500
    assert retrieve_step["chunk_overlap"] == 50


def test_settings_chunking_preset_mode_ignores_manual_values_and_applies(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "preset-mode.txt",
                ("Preset mode token. " * 90).encode("utf-8"),
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 201

    applied = client.post(
        "/api/v1/settings/chunking",
        json={
            "mode": "small",
            "chunk_size": 99999,  # must be ignored in preset mode
            "chunk_overlap": -5,  # must be ignored in preset mode
        },
    )
    assert applied.status_code == 200
    body = applied.json()
    assert body["status"] == "reindexed"
    assert body["mode"] == "small"
    assert body["chunk_mode"] == "preset"
    assert body["chunk_size"] == 500
    assert body["chunk_overlap"] == 50
    assert body["reindexed_documents"] >= 1

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "Preset mode token là gì?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_response.status_code == 200
    retrieve_step = query_response.json()["trace"][0]
    assert retrieve_step["step"] == "retrieve"
    assert retrieve_step["chunk_size"] == 500
    assert retrieve_step["chunk_overlap"] == 50


def test_settings_chunking_custom_mode_accepts_and_applies_values(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "custom-mode.txt",
                ("Custom mode token. " * 80).encode("utf-8"),
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 201

    applied = client.post(
        "/api/v1/settings/chunking",
        json={
            "mode": "custom",
            "chunk_size": 700,
            "chunk_overlap": 120,
        },
    )
    assert applied.status_code == 200
    body = applied.json()
    assert body["status"] == "reindexed"
    assert body["mode"] == "custom"
    assert body["chunk_mode"] == "custom"
    assert body["chunk_size"] == 700
    assert body["chunk_overlap"] == 120
    assert body["reindexed_documents"] >= 1

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "Custom mode token là gì?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_response.status_code == 200
    retrieve_step = query_response.json()["trace"][0]
    assert retrieve_step["chunk_size"] == 700
    assert retrieve_step["chunk_overlap"] == 120


def test_settings_chunking_custom_mode_rejects_invalid_payload(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    missing_values = client.post(
        "/api/v1/settings/chunking",
        json={"mode": "custom"},
    )
    assert missing_values.status_code == 422

    invalid_overlap = client.post(
        "/api/v1/settings/chunking",
        json={
            "mode": "custom",
            "chunk_size": 500,
            "chunk_overlap": 500,
        },
    )
    assert invalid_overlap.status_code == 422


def test_settings_chunking_change_triggers_reindex_and_updates_chunk_count(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "reindex-check.txt",
                ("Reindex chunk-count token. " * 220).encode("utf-8"),
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 201

    small = client.post("/api/v1/settings/chunking", json={"mode": "small"})
    assert small.status_code == 200
    small_body = small.json()
    assert small_body["chunk_size"] == 500
    assert small_body["chunk_overlap"] == 50
    assert small_body["active_chunks"] > 0

    large = client.post("/api/v1/settings/chunking", json={"mode": "large"})
    assert large.status_code == 200
    large_body = large.json()
    assert large_body["chunk_size"] == 1500
    assert large_body["chunk_overlap"] == 200
    assert large_body["active_chunks"] > 0
    assert large_body["active_chunks"] < small_body["active_chunks"]


def test_settings_retrieval_preset_modes_map_correctly(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    low = client.post(
        "/api/v1/settings/retrieval",
        json={"mode": "low", "top_k": 20},
    )
    assert low.status_code == 200
    low_body = low.json()
    assert low_body["status"] == "updated"
    assert low_body["mode"] == "low"
    assert low_body["retrieval_mode"] == "preset"
    assert low_body["top_k"] == 3
    assert low_body["rerank_top_n"] <= low_body["top_k"]

    balanced = client.post("/api/v1/settings/retrieval", json={"mode": "balanced"})
    assert balanced.status_code == 200
    assert balanced.json()["top_k"] == 5

    high = client.post("/api/v1/settings/retrieval", json={"mode": "high"})
    assert high.status_code == 200
    assert high.json()["top_k"] == 8


def test_settings_retrieval_custom_top_k_updates_query_retrieval(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "retrieval-custom.txt",
                ("Retrieval custom token alpha beta gamma. " * 220).encode("utf-8"),
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 201

    applied = client.post(
        "/api/v1/settings/retrieval",
        json={"mode": "custom", "top_k": 4},
    )
    assert applied.status_code == 200
    body = applied.json()
    assert body["mode"] == "custom"
    assert body["retrieval_mode"] == "custom"
    assert body["top_k"] == 4
    assert body["rerank_top_n"] <= body["top_k"]

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "Retrieval custom token là gì?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_response.status_code == 200
    query_body = query_response.json()
    retrieve_step = query_body["trace"][0]
    rerank_step = query_body["trace"][1]
    context_step = query_body["trace"][2]
    assert retrieve_step["step"] == "retrieve"
    assert retrieve_step["top_k"] == 4
    assert retrieve_step["count"] <= 4
    assert rerank_step["rerank_top_n"] <= retrieve_step["top_k"]
    assert context_step["final_context_size"] == context_step["count"]


def test_settings_retrieval_invalid_custom_top_k_rejected(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    missing_top_k = client.post(
        "/api/v1/settings/retrieval",
        json={"mode": "custom"},
    )
    assert missing_top_k.status_code == 422

    too_small = client.post(
        "/api/v1/settings/retrieval",
        json={"mode": "custom", "top_k": 0},
    )
    assert too_small.status_code == 422

    too_large = client.post(
        "/api/v1/settings/retrieval",
        json={"mode": "custom", "top_k": 21},
    )
    assert too_large.status_code == 422


def test_query_falls_back_to_seeded_after_delete_all_uploaded_documents(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_response = client.post(
        "/api/v1/documents/upload",
        files={
            "file": (
                "to-delete.txt",
                b"uploaded-doc-token that will be removed before query",
                "text/plain",
            )
        },
    )
    assert upload_response.status_code == 201

    deleted = client.delete("/api/v1/documents")
    assert deleted.status_code == 200
    assert deleted.json()["deleted_documents"] == 1

    query_response = client.post(
        "/api/v1/query",
        json={
            "query": "What is in the seeded corpus?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_response.status_code == 200
    body = query_response.json()
    assert body["trace"][0]["index_source"] == "seeded"


def test_uploaded_retrieval_excludes_deleted_document_chunks(
    isolated_client: tuple[TestClient, Path],
) -> None:
    client, _ = isolated_client

    upload_a = client.post(
        "/api/v1/documents/upload",
        files={"file": ("doc-a.txt", b"alpha-doc-token-44221", "text/plain")},
    )
    upload_b = client.post(
        "/api/v1/documents/upload",
        files={"file": ("doc-b.txt", b"beta-doc-token-99117", "text/plain")},
    )
    assert upload_a.status_code == 201
    assert upload_b.status_code == 201
    doc_a = upload_a.json()["document_id"]
    doc_b = upload_b.json()["document_id"]

    query_a_before_delete = client.post(
        "/api/v1/query",
        json={
            "query": "alpha-doc-token-44221 nằm ở đâu?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_a_before_delete.status_code == 200
    assert any(citation["doc_id"] == doc_a for citation in query_a_before_delete.json()["citations"])

    query_b = client.post(
        "/api/v1/query",
        json={
            "query": "beta-doc-token-99117 có trong tài liệu nào?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_b.status_code == 200
    assert any(citation["doc_id"] == doc_b for citation in query_b.json()["citations"])

    deleted = client.delete(f"/api/v1/documents/{doc_a}")
    assert deleted.status_code == 200

    query_after_delete = client.post(
        "/api/v1/query",
        json={
            "query": "alpha-doc-token-44221 nằm ở đâu?",
            "mode": "standard",
            "chat_history": [],
        },
    )
    assert query_after_delete.status_code == 200
    body = query_after_delete.json()
    assert body["trace"][0]["index_source"] == "uploaded"
    assert all(citation["doc_id"] != doc_a for citation in body["citations"])


def test_restart_with_empty_registry_does_not_use_stale_uploaded_indexes(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    data_dir = tmp_path / "data"
    corpus_dir = tmp_path / "corpus"
    corpus_dir.mkdir(parents=True, exist_ok=True)
    (corpus_dir / "seed.txt").write_text(
        "seeded-token-only-7781 should be used after stale uploaded index cleanup.",
        encoding="utf-8",
    )

    monkeypatch.setenv("DATA_DIR", str(data_dir))
    monkeypatch.setenv("INDEX_DIR", str(data_dir / "indexes"))
    monkeypatch.setenv("CORPUS_DIR", str(corpus_dir))
    get_settings.cache_clear()

    app_first = create_app()
    client_first = TestClient(app_first)
    try:
        upload = client_first.post(
            "/api/v1/documents/upload",
            files={"file": ("stale.txt", b"stale-uploaded-token-5566", "text/plain")},
        )
        assert upload.status_code == 201
        stale_doc_id = upload.json()["document_id"]
    finally:
        client_first.close()

    assert (data_dir / "indexes" / "vector_index.json").exists()
    assert (data_dir / "indexes" / "bm25_index.json").exists()

    (data_dir / "document_registry.json").write_text('{"documents": []}', encoding="utf-8")
    raw_dir = data_dir / "raw"
    for path in raw_dir.rglob("*"):
        if path.is_file():
            path.unlink()

    get_settings.cache_clear()
    app_second = create_app()
    client_second = TestClient(app_second)
    try:
        stale_query = client_second.post(
            "/api/v1/query",
            json={
                "query": "stale-uploaded-token-5566",
                "mode": "standard",
                "chat_history": [],
            },
        )
        assert stale_query.status_code == 200
        stale_body = stale_query.json()
        assert stale_body["trace"][0]["index_source"] == "seeded"
        assert all(citation["doc_id"] != stale_doc_id for citation in stale_body["citations"])

        seeded_query = client_second.post(
            "/api/v1/query",
            json={
                "query": "seeded-token-only-7781 là gì?",
                "mode": "standard",
                "chat_history": [],
            },
        )
        assert seeded_query.status_code == 200
        seeded_body = seeded_query.json()
        assert seeded_body["trace"][0]["index_source"] == "seeded"
    finally:
        client_second.close()
        get_settings.cache_clear()
