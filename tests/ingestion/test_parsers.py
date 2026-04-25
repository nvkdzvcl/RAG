"""Tests for parser abstraction and mixed-content block extraction."""

from __future__ import annotations

from pathlib import Path
from types import SimpleNamespace

from app.ingestion.chunker import Chunker
from app.ingestion.parsers.docx_parser import DocxParser
from app.ingestion.parsers.pdf_parser import PDFParser
from app.ingestion.parsers.text_parser import TextParser
from app.ingestion.text_loader import TextLoader


def test_pdf_parser_extracts_text_table_and_image_blocks(monkeypatch) -> None:
    class FakePage:
        images = [{"name": "img_1", "x0": 1, "top": 2, "x1": 3, "bottom": 4}]

        @staticmethod
        def extract_text() -> str:
            return "Muc Luc\n\nDoan van thu nhat.\n\nDoan van thu hai."

        @staticmethod
        def extract_tables() -> list[list[list[str]]]:
            return [[["Name", "Score"], ["A", "10"]]]

    class FakePDF:
        def __init__(self) -> None:
            self.pages = [FakePage()]

        def __enter__(self) -> "FakePDF":
            return self

        def __exit__(self, exc_type, exc, tb) -> None:
            _ = exc_type
            _ = exc
            _ = tb

    monkeypatch.setattr(
        "app.ingestion.parsers.pdf_parser.pdfplumber",
        SimpleNamespace(open=lambda _: FakePDF()),
    )

    parser = PDFParser()
    blocks = parser.parse(Path("sample.pdf"))

    assert any(block.type == "text" for block in blocks)
    assert any(block.type == "table" for block in blocks)
    assert any(block.type == "image" for block in blocks)
    assert all("page" in block.metadata for block in blocks)


def test_docx_parser_extracts_heading_paragraph_and_table(tmp_path: Path) -> None:
    from docx import Document

    doc_path = tmp_path / "mixed.docx"
    doc = Document()
    doc.add_heading("Bao cao", level=1)
    doc.add_paragraph("Noi dung chinh cua tai lieu.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Score"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "10"
    doc.save(str(doc_path))

    parser = DocxParser()
    blocks = parser.parse(doc_path)

    assert any(block.type == "text" and block.metadata.get("is_heading") for block in blocks)
    assert any(block.type == "table" for block in blocks)
    table_block = next(block for block in blocks if block.type == "table")
    assert "| Name | Score |" in table_block.content


def test_vietnamese_text_preserved_through_loader_and_chunker(tmp_path: Path) -> None:
    path = tmp_path / "vi.txt"
    vi_text = "Tiếng Việt có dấu: Trường đại học, nghiên cứu, dữ liệu."
    path.write_text(vi_text, encoding="utf-8")

    docs = TextLoader().load(path, metadata={"lang": "vi"})
    chunks = Chunker(chunk_size=16, chunk_overlap=2).chunk_documents(docs)

    assert docs[0].content == vi_text
    assert any("Tiếng Việt có dấu" in chunk.content for chunk in chunks)
    assert all(chunk.metadata.get("language") == "auto" or chunk.metadata.get("language") == "vi" for chunk in chunks)
    assert all(chunk.metadata.get("block_type") == "text" for chunk in chunks)


def test_text_parser_splits_paragraphs_without_losing_utf8(tmp_path: Path) -> None:
    path = tmp_path / "mixed.txt"
    path.write_text("Đoạn một.\n\nĐoạn hai có bảng giả lập.", encoding="utf-8")

    blocks = TextParser().parse(path)

    assert len(blocks) == 2
    assert blocks[0].content == "Đoạn một."
    assert "Đoạn hai" in blocks[1].content


def test_pdf_parser_ocr_disabled_does_not_invoke_ocr(monkeypatch) -> None:
    class FakePage:
        images = []

        @staticmethod
        def extract_text() -> str:
            return ""

        @staticmethod
        def extract_tables() -> list[list[list[str]]]:
            return []

    class FakePDF:
        def __init__(self) -> None:
            self.pages = [FakePage()]

        def __enter__(self) -> "FakePDF":
            return self

        def __exit__(self, exc_type, exc, tb) -> None:
            _ = exc_type
            _ = exc
            _ = tb

    monkeypatch.setattr(
        "app.ingestion.parsers.pdf_parser.pdfplumber",
        SimpleNamespace(open=lambda _: FakePDF()),
    )

    def _unexpected_ocr(*args, **kwargs) -> str:
        _ = args
        _ = kwargs
        raise AssertionError("OCR should not run when OCR is disabled")

    monkeypatch.setattr("app.ingestion.parsers.pdf_parser.ocr_pdf_page_with_pymupdf", _unexpected_ocr)

    parser = PDFParser(ocr_enabled=False, ocr_min_text_chars=100)
    blocks = parser.parse(Path("sample.pdf"))

    assert blocks == []


def test_pdf_parser_missing_tesseract_does_not_crash(monkeypatch) -> None:
    class FakePage:
        images = []

        @staticmethod
        def extract_text() -> str:
            return ""

        @staticmethod
        def extract_tables() -> list[list[list[str]]]:
            return []

    class FakePDF:
        def __init__(self) -> None:
            self.pages = [FakePage()]

        def __enter__(self) -> "FakePDF":
            return self

        def __exit__(self, exc_type, exc, tb) -> None:
            _ = exc_type
            _ = exc
            _ = tb

    monkeypatch.setattr(
        "app.ingestion.parsers.pdf_parser.pdfplumber",
        SimpleNamespace(open=lambda _: FakePDF()),
    )
    monkeypatch.setattr("app.ingestion.parsers.pdf_parser.is_tesseract_available", lambda: False)

    parser = PDFParser(ocr_enabled=True, ocr_min_text_chars=100)
    blocks = parser.parse(Path("scan.pdf"))

    assert blocks == []


def test_pdf_parser_appends_ocr_text_block_with_metadata(monkeypatch) -> None:
    class FakePage:
        images = []

        @staticmethod
        def extract_text() -> str:
            return ""

        @staticmethod
        def extract_tables() -> list[list[list[str]]]:
            return []

    class FakePDF:
        def __init__(self) -> None:
            self.pages = [FakePage()]

        def __enter__(self) -> "FakePDF":
            return self

        def __exit__(self, exc_type, exc, tb) -> None:
            _ = exc_type
            _ = exc
            _ = tb

    monkeypatch.setattr(
        "app.ingestion.parsers.pdf_parser.pdfplumber",
        SimpleNamespace(open=lambda _: FakePDF()),
    )
    monkeypatch.setattr("app.ingestion.parsers.pdf_parser.is_tesseract_available", lambda: True)
    ocr_calls: list[tuple[tuple[object, ...], dict[str, object]]] = []

    def _fake_ocr(*args, **kwargs) -> str:
        ocr_calls.append((args, kwargs))
        return "Nội dung OCR tiếng Việt có dấu."

    monkeypatch.setattr(
        "app.ingestion.parsers.pdf_parser.ocr_pdf_page_with_pymupdf",
        _fake_ocr,
    )

    parser = PDFParser(
        ocr_enabled=True,
        ocr_language="vie+eng",
        ocr_min_text_chars=100,
        ocr_render_dpi=216,
        ocr_confidence_threshold=40.0,
    )
    blocks = parser.parse(Path("scan.pdf"))

    assert len(blocks) == 1
    block = blocks[0]
    assert block.type == "text"
    assert "OCR tiếng Việt" in block.content
    assert block.metadata["ocr"] is True
    assert block.metadata["block_type"] == "ocr_text"
    assert block.metadata["ocr_language"] == "vie+eng"
    assert block.metadata["language"] == "vi"
    assert len(ocr_calls) == 1


def test_pdf_parser_text_extraction_still_works_without_ocr_fallback(monkeypatch) -> None:
    class FakePage:
        images = []

        @staticmethod
        def extract_text() -> str:
            return "This page already has enough digital text to skip OCR fallback."

        @staticmethod
        def extract_tables() -> list[list[list[str]]]:
            return []

    class FakePDF:
        def __init__(self) -> None:
            self.pages = [FakePage()]

        def __enter__(self) -> "FakePDF":
            return self

        def __exit__(self, exc_type, exc, tb) -> None:
            _ = exc_type
            _ = exc
            _ = tb

    monkeypatch.setattr(
        "app.ingestion.parsers.pdf_parser.pdfplumber",
        SimpleNamespace(open=lambda _: FakePDF()),
    )
    monkeypatch.setattr("app.ingestion.parsers.pdf_parser.is_tesseract_available", lambda: True)

    def _unexpected_ocr(*args, **kwargs) -> str:
        _ = args
        _ = kwargs
        raise AssertionError("OCR should be skipped when page already has enough text")

    monkeypatch.setattr("app.ingestion.parsers.pdf_parser.ocr_pdf_page_with_pymupdf", _unexpected_ocr)

    parser = PDFParser(ocr_enabled=True, ocr_min_text_chars=10)
    blocks = parser.parse(Path("digital.pdf"))

    assert any(block.type == "text" for block in blocks)
    assert all(block.metadata.get("block_type") != "ocr_text" for block in blocks)


def test_pdf_parser_ocr_runtime_error_does_not_crash(monkeypatch) -> None:
    class FakePage:
        images = []

        @staticmethod
        def extract_text() -> str:
            return ""

        @staticmethod
        def extract_tables() -> list[list[list[str]]]:
            return []

    class FakePDF:
        def __init__(self) -> None:
            self.pages = [FakePage()]

        def __enter__(self) -> "FakePDF":
            return self

        def __exit__(self, exc_type, exc, tb) -> None:
            _ = exc_type
            _ = exc
            _ = tb

    monkeypatch.setattr(
        "app.ingestion.parsers.pdf_parser.pdfplumber",
        SimpleNamespace(open=lambda _: FakePDF()),
    )
    monkeypatch.setattr("app.ingestion.parsers.pdf_parser.is_tesseract_available", lambda: True)

    def _broken_ocr(*args, **kwargs) -> str:
        _ = args
        _ = kwargs
        raise RuntimeError("missing OCR dependency")

    monkeypatch.setattr("app.ingestion.parsers.pdf_parser.ocr_pdf_page_with_pymupdf", _broken_ocr)

    parser = PDFParser(ocr_enabled=True, ocr_min_text_chars=100)
    blocks = parser.parse(Path("broken-ocr.pdf"))

    assert blocks == []
