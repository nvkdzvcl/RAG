"""Tests for ingestion loaders output shape and metadata handling."""

from pathlib import Path
from types import SimpleNamespace

from app.ingestion.chunker import Chunker
from app.ingestion.docx_loader import DocxLoader
from app.ingestion.markdown_loader import MarkdownLoader
from app.ingestion.pdf_loader import PdfLoader
from app.ingestion.parsers.pdf_parser import PDFParser
from app.ingestion.text_loader import TextLoader


def test_text_loader_output_shape(tmp_path: Path) -> None:
    source = tmp_path / "sample.txt"
    source.write_text("Line 1\nLine 2\n", encoding="utf-8")

    loader = TextLoader()
    docs = loader.load(source, metadata={"source_type": "notes"})

    assert len(docs) == 1
    doc = docs[0]
    assert doc.doc_id
    assert doc.source == str(source)
    assert doc.title == "sample"
    assert doc.section is None
    assert doc.page is None
    assert doc.content.startswith("Line 1")
    assert doc.block_type == "text"
    assert doc.language == "auto"
    assert doc.metadata["source_type"] == "notes"


def test_markdown_loader_preserves_section_metadata(tmp_path: Path) -> None:
    source = tmp_path / "guide.md"
    source.write_text(
        "# Getting Started\n\nOverview text.\n\n## Setup\n\nSetup steps.",
        encoding="utf-8",
    )

    loader = MarkdownLoader()
    docs = loader.load(source, metadata={"lang": "en"})

    assert len(docs) >= 4
    assert all(doc.doc_id == docs[0].doc_id for doc in docs)
    assert all(doc.source == str(source) for doc in docs)
    assert all(doc.title == "Getting Started" for doc in docs)
    assert any(doc.section == "Getting Started" for doc in docs)
    assert any(doc.section == "Setup" for doc in docs)
    assert all(doc.page is None for doc in docs)
    assert all(doc.block_type in {"text", "table", "image"} for doc in docs)
    assert all(doc.metadata["lang"] == "en" for doc in docs)


def test_docx_loader_supports_vietnamese_text(tmp_path: Path) -> None:
    doc_path = tmp_path / "vietnamese.docx"

    from docx import Document

    doc = Document()
    doc.add_heading("Giới thiệu", level=1)
    doc.add_paragraph("Hệ thống hỗ trợ tiếng Việt đầy đủ dấu.")
    doc.save(str(doc_path))

    loader = DocxLoader()
    docs = loader.load(doc_path, metadata={"lang": "vi"})

    assert docs
    assert any("tiếng Việt đầy đủ dấu" in doc.content for doc in docs)
    assert all(doc.metadata["lang"] == "vi" for doc in docs)


def test_pdf_loader_ocr_vietnamese_text_becomes_chunk_and_preserves_metadata(
    tmp_path: Path,
    monkeypatch,
) -> None:
    pdf_path = tmp_path / "scan.pdf"
    pdf_path.write_bytes(b"%PDF-1.4 fake content")

    class FakePage:
        images: list[object] = []

        @staticmethod
        def extract_text() -> str:
            return ""

        @staticmethod
        def extract_tables() -> list[list[list[str]]]:
            return []

    class FakePDF:
        def __init__(self) -> None:
            self.pages = [FakePage()]

        def __enter__(self) -> "FakePDF":
            return self

        def __exit__(self, exc_type, exc, tb) -> None:
            _ = exc_type
            _ = exc
            _ = tb

    monkeypatch.setattr(
        "app.ingestion.parsers.pdf_parser.pdfplumber",
        SimpleNamespace(open=lambda _: FakePDF()),
    )
    monkeypatch.setattr(
        "app.ingestion.parsers.pdf_parser.is_tesseract_available", lambda: True
    )
    monkeypatch.setattr(
        "app.ingestion.parsers.pdf_parser.ocr_pdf_page_with_pymupdf",
        lambda *args, **kwargs: (
            "Tài liệu quét có nội dung tiếng Việt: dữ liệu truy hồi."
        ),
    )

    loader = PdfLoader(
        parser=PDFParser(
            ocr_enabled=True,
            ocr_language="vie+eng",
            ocr_min_text_chars=100,
            ocr_render_dpi=216,
            ocr_confidence_threshold=40.0,
        )
    )
    docs = loader.load(pdf_path, metadata={"lang": "vi"})

    assert docs
    assert any("tiếng Việt" in doc.content for doc in docs)
    assert any(doc.metadata.get("ocr") is True for doc in docs)
    assert any(doc.metadata.get("block_type") == "ocr_text" for doc in docs)

    chunks = Chunker(chunk_size=40, chunk_overlap=5).chunk_documents(docs)
    assert chunks
    assert any("tiếng Việt" in chunk.content for chunk in chunks)
    assert any(chunk.metadata.get("block_type") == "ocr_text" for chunk in chunks)
